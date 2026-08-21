# Object structure:
# DOCX:  document -> document.paragraphs -> paragraph -> paragraph.text
# PDF:   document -> page                -> page      -> page.get_text()
# XLSX:  document -> sheet     -> row    -> cell      -> cell.value
# Image: image    -> pillow image        -> OCR       -> pytesseract.image_to_string

from pathlib import Path
from docx import Document
from openpyxl import load_workbook
from PIL import Image
import pymupdf
import xlrd
import pytesseract # Tesseract OCR
from jobmatch.document.models import DocumentContent


def ocr_image(image):
    text = pytesseract.image_to_string(
        image,
        lang="eng+jpn+chi_sim+chi_tra"
    )

    return text.strip() # Remove leading/trailing whitespace


def extract_docx(file_path):
    path = Path(file_path)
    text = []

    document = Document(path)

    # DOCX structure:
    # document            -> the entire Word document
    # document.paragraphs -> a list of paragraph objects in the document
    # paragraph           -> one paragraph object
    # paragraph.text      -> the text (str) inside that paragraph

    for paragraph in document.paragraphs:
        text.append(paragraph.text)

    return DocumentContent(
        source_file=path.name,
        file_type="docx",
        text="\n".join(text)
    )


def extract_pdf(file_path):
    path = Path(file_path)
    text = []

    document = pymupdf.open(path)

    for page in document:
        page_text = page.get_text().strip()

        # case 1: If the page has text, use it directly
        if page_text:
            text.append(page_text)

        # case 2: If the page has no text, use OCR to extract text from the image
        else:
            page_image = page.get_pixmap() # Get the pixelmap of the page

            image = Image.frombytes(
                "RGB",
                [page_image.width, page_image.height],
                page_image.samples
            )

            page_text = ocr_image(image)
            text.append(page_text)

    document.close()

    return DocumentContent(
        source_file=path.name,
        file_type="pdf",
        text="\n".join(text)
    )


def extract_xlsx(file_path):
    path = Path(file_path)
    text = []

    document = load_workbook(path, data_only=True)

    for worksheet in document.worksheets:
        for row in worksheet.iter_rows():
            row_text = []

            for cell in row:
                if cell.value is not None:
                    row_text.append(str(cell.value))

            if row_text:
                text.append(" | ".join(row_text))

    document.close()

    return DocumentContent(
        source_file=path.name,
        file_type="xlsx",
        text="\n".join(text)
    )


def extract_xls(file_path):
    path = Path(file_path)
    text = []

    document = xlrd.open_workbook(path)

    for worksheet in document.sheets():
        for row_index in range(worksheet.nrows):
            row_text = []

            for cell in worksheet.row(row_index):
                if cell.value != "":
                    row_text.append(str(cell.value))

            if row_text:
                text.append(" | ".join(row_text))

    document.release_resources()

    return DocumentContent(
        source_file=path.name,
        file_type="xls",
        text="\n".join(text)
    )


def extract_image(file_path):
    path = Path(file_path)

    image = Image.open(path)
    text = ocr_image(image)
    image.close()

    return DocumentContent(
        source_file=path.name,
        # ".jpg", ".jpeg", ".png" -> "jpg", "jpeg", "png"
        file_type=path.suffix.lower().lstrip("."),
        text=text
    )