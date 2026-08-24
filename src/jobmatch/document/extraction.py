# Object structure:
# DOCX:  document -> document.paragraphs -> paragraph -> paragraph.text
# PDF:   document -> page                -> page      -> page.get_text()
# XLSX:  document -> sheet     -> row    -> cell      -> cell.value
# Image: image    -> pillow image        -> OCR       -> pytesseract.image_to_string

# DOCX / PDF (layout-aware):
# file -> Docling DocumentConverter -> DoclingDocument -> Markdown


from pathlib import Path
from docx import Document
from openpyxl import load_workbook
from PIL import Image
from docling.document_converter import DocumentConverter
import pymupdf
import xlrd
import pytesseract # Tesseract OCR

from jobmatch.document.models import DocumentContent


def ocr_image(image):
    text = pytesseract.image_to_string(
        image,
        lang="eng+jpn+chi_sim+chi_tra"
    )

    return text.strip() # Remove leading/trailing whitespace


converter = DocumentConverter()
def extract_docling(file_path):
    """
    Extract document content using Docling.

    Docling performs layout-aware document parsing and exports the result
    as Markdown so document structure can be preserved better than plain
    text extraction.
    """

    # converter = DocumentConverter()
    result = converter.convert(str(file_path))

    return result.document.export_to_markdown()


def extract_docx(file_path):
    path = Path(file_path)

    # Primary method:
    # Use Docling to preserve document structure such as headings, lists,
    # tables, and more complex layouts.
    try:
        print("DOCX extraction path: Docling")
        text = extract_docling(path)

        return DocumentContent(
            source_file=path.name,
            file_type="docx",
            text=text
        )

    except Exception:
        # Fallback method:
        # If Docling cannot process the document, use python-docx to extract
        # standard paragraph text.
        print("DOCX extraction path: python-docx fallback")
        text = []

        document = Document(path)

        # DOCX structure:
        # document            -> the entire Word document
        # document.paragraphs -> a list of paragraph objects in the document
        # paragraph           -> one paragraph object
        # paragraph.text      -> the text (str) inside that paragraph
        #
        # Limitation:
        # This basic method may not preserve complex layouts such as
        # multi-column content, text boxes, or some table structures.

        for paragraph in document.paragraphs:
            text.append(paragraph.text)

        return DocumentContent(
            source_file=path.name,
            file_type="docx",
            text="\n".join(text)
        )


def extract_pdf(file_path):
    path = Path(file_path)

    # Primary method:
    # Use Docling for layout-aware PDF parsing.
    try:
        print("PDF extraction path: Docling")
        text = extract_docling(path)

        return DocumentContent(
            source_file=path.name,
            file_type="pdf",
            text=text
        )

    except Exception:
        # Fallback method:
        # If Docling fails, use PyMuPDF for normal text PDFs and Tesseract OCR
        # for pages where no embedded text can be extracted.
        print("PDF extraction path: PyMuPDF/Tesseract fallback")
        text = []

        document = pymupdf.open(path)

        for page in document:
            page_text = page.get_text().strip()

            # Case 1:
            # If the page has embedded text, use it directly.
            if page_text:
                text.append(page_text)

            # Case 2:
            # If the page has no embedded text, render it as an image and
            # use OCR to extract the text.
            else:
                page_image = page.get_pixmap()  # Get the pixmap of the page

                image = Image.frombytes(
                    "RGB",
                    [page_image.width, page_image.height],
                    page_image.samples
                )

                page_text = ocr_image(image)
                text.append(page_text)

        document.close()

        return DocumentContent(
            source_file=path.name,
            file_type="pdf",
            text="\n".join(text)
        )



def extract_xlsx(file_path):
    path = Path(file_path)
    text = []

    document = load_workbook(path, data_only=True)

    for worksheet in document.worksheets:
        for row in worksheet.iter_rows():
            row_text = []

            for cell in row:
                if cell.value is not None:
                    row_text.append(str(cell.value))

            if row_text:
                text.append(" | ".join(row_text))

    document.close()

    return DocumentContent(
        source_file=path.name,
        file_type="xlsx",
        text="\n".join(text)
    )


def extract_xls(file_path):
    path = Path(file_path)
    text = []

    document = xlrd.open_workbook(path)

    for worksheet in document.sheets():
        for row_index in range(worksheet.nrows):
            row_text = []

            for cell in worksheet.row(row_index):
                if cell.value != "":
                    row_text.append(str(cell.value))

            if row_text:
                text.append(" | ".join(row_text))

    document.release_resources()

    return DocumentContent(
        source_file=path.name,
        file_type="xls",
        text="\n".join(text)
    )


def extract_image(file_path):
    path = Path(file_path)

    image = Image.open(path)
    text = ocr_image(image)
    image.close()

    return DocumentContent(
        source_file=path.name,
        # ".jpg", ".jpeg", ".png" -> "jpg", "jpeg", "png"
        file_type=path.suffix.lower().lstrip("."),
        text=text
    )